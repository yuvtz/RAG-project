#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
index_documents.py
Index PDF/DOCX into Postgres (pgvector) using Google Gemini embeddings.

Examples:
  py index_documents.py --paths docs/ --strategy sentences     # each sentence/heading/bullet = a separate chunk
  py index_documents.py --paths docs/ --strategy fixed --chunk-size 800 --overlap 120
"""

import os
import re
import glob
import argparse
import time
import string
from typing import List, Tuple

from dotenv import load_dotenv
import google.generativeai as genai
import psycopg2
from psycopg2.extras import execute_values
from pypdf import PdfReader
from docx import Document
from tqdm import tqdm


# ------------------------------ File IO ------------------------------

def read_pdf(path: str) -> str:
    """Read a PDF and return concatenated text of all pages."""
    reader = PdfReader(path)
    parts = []
    for p in reader.pages:
        parts.append(p.extract_text() or "")
    return "\n".join(parts)

def read_docx(path: str) -> str:
    """
    Structured DOCX reading.
    - Headings are preserved as standalone lines (Heading 1/2/3…)
    - Lists get a leading dash if no bullet/numbering is present
    - Table rows are collected as new lines
    """
    doc = Document(path)
    lines: List[str] = []

    # Paragraphs (including headings/lists)
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        style = p.style.name if p.style is not None else ""

        if t == "":
            lines.append("")  # paragraph boundary
            continue

        # Heading
        if "Heading" in style:
            lines.append(t)
            lines.append("")  # blank line after heading
            continue

        # Lists
        if "List" in style or "Bullet" in style or "Number" in style:
            if re.match(r'^\s*(?:[-*•–]|[0-9]+[.)])\s+', t):
                lines.append(t)
            else:
                lines.append(f"- {t}")
            continue

        # Regular paragraph
        lines.append(t)

    # Tables → one output line per table row
    for tbl in doc.tables:
        for row in tbl.rows:
            cell_text = " | ".join((cell.text or "").strip() for cell in row.cells)
            if cell_text:
                lines.append(cell_text)

    return "\n".join(lines)

def read_any(path: str) -> str:
    """Read a file by extension (PDF/DOCX); raise on unsupported types."""
    ext = os.path.splitext(path.lower())[1]
    if ext == ".pdf":
        return read_pdf(path)
    if ext in (".docx", ".doc"):
        return read_docx(path)
    raise ValueError(f"Unsupported file type: {ext}")

def discover_files(paths: List[str]) -> List[str]:
    """Expand directories/globs into a unique sorted list of PDF/DOCX files."""
    found = []
    for p in paths:
        if os.path.isdir(p):
            found += glob.glob(os.path.join(p, "**", "*.pdf"), recursive=True)
            found += glob.glob(os.path.join(p, "**", "*.docx"), recursive=True)
        elif os.path.isfile(p):
            found.append(p)
        else:
            found += glob.glob(p)
    return sorted(set(found))


# ------------------------------ Cleaning ------------------------------

_WS_RE = re.compile(r"[ \t]+")

def clean_text(text: str) -> str:
    """Normalize whitespace and collapse excessive blank lines."""
    text = text.replace("\u00A0", " ")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = _WS_RE.sub(" ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ------------------------------ Splitting ------------------------------

_BULLET_RE = re.compile(r'^\s*(?:[-*•–]|[0-9]+[.)])\s+')
_END_PUNCT_RE = re.compile(r'[.!?…]$')

def _looks_like_heading(line: str) -> bool:
    """Heuristic heading: short, no terminal punctuation, ALL-CAPS/Title-Case or few words."""
    l = line.strip()
    if not l or _END_PUNCT_RE.search(l):
        return False
    short = len(l) <= 80
    capsish = (l.isupper() and any(ch in string.ascii_letters for ch in l)) or l.istitle()
    return short and (capsish or len(l.split()) <= 8)

def _segment_units(text: str, lang: str = "en") -> List[str]:
    """
    Split into readable units independent of source (PDF/DOCX):
    - Headings as standalone lines
    - Bullets/numbered items as standalone units (append period if missing)
    - Paragraphs → sentences (pysbd preferred; regex fallback)
    """
    lines = [ln.rstrip() for ln in text.splitlines()]
    units: List[str] = []
    para_buf: List[str] = []

    def flush_para():
        """Flush accumulated paragraph buffer into sentence units."""
        nonlocal para_buf, units
        if not para_buf:
            return
        para = " ".join(s for s in para_buf if s.strip())
        para_buf = []
        if not para.strip():
            return
        try:
            import pysbd  # type: ignore
            seg = pysbd.Segmenter(language=lang, clean=True)
            sents = [s.strip() for s in seg.segment(para) if s.strip()]
        except Exception:
            sents = [s.strip() for s in re.split(r'(?<=[.!?])\s+', para) if s.strip()]
        units.extend(sents)

    for raw in lines:
        ln = raw.strip()
        if ln == "":
            flush_para()
            continue
        if _BULLET_RE.match(ln):
            flush_para()
            units.append(ln if _END_PUNCT_RE.search(ln) else ln + ".")
            continue
        if _looks_like_heading(ln):
            flush_para()
            units.append(ln)
            continue
        para_buf.append(ln)

    flush_para()
    return units

def split_fixed(text: str, chunk_size: int = 800, overlap: int = 120) -> List[str]:
    """Fixed-size windowing with overlap."""
    if chunk_size <= 0:
        raise ValueError("chunk_size must be positive")
    if overlap < 0:
        raise ValueError("overlap must be >= 0")
    out, i, n = [], 0, len(text)
    while i < n:
        j = min(i + chunk_size, n)
        out.append(text[i:j])
        if j == n:
            break
        i = max(0, j - overlap)
    return out

def split_sentences(text: str, max_chars: int = 900, lang: str = "en", merge: bool = True) -> List[str]:
    """
    Smart splitting into sentences/headings/bullets.
    Default merges; do_split calls with merge=False to disable merging.
    """
    atoms = _segment_units(text, lang=lang)
    if not merge:
        return atoms  # sentence/heading/bullet = single chunk
    # (Not used currently but kept for future options)
    chunks, cur = [], ""
    for a in atoms:
        if not cur:
            cur = a
        elif len(cur) + 1 + len(a) <= max_chars:
            cur = f"{cur} {a}"
        else:
            chunks.append(cur)
            cur = a
    if cur:
        chunks.append(cur)
    return chunks

def split_paragraphs(text: str) -> List[str]:
    """
    Paragraph-first splitting based on document structure:
    - Headings and bullets become standalone chunks.
    - Consecutive sentences accumulate into a paragraph until a new heading/bullet.
    - Very long paragraphs are softly split by sentences (internal safety cap).
    """
    # Base units: headings / bullets / sentences
    atoms = _segment_units(text, lang="en")

    def _is_heading(s: str) -> bool:
        return _looks_like_heading(s)

    def _is_bullet(s: str) -> bool:
        return _BULLET_RE.match(s) is not None

    out: List[str] = []
    cur: List[str] = []

    def flush_cur():
        """Push current paragraph; softly split if extremely long."""
        if not cur:
            return
        paragraph = " ".join(cur).strip()
        cur.clear()
        if not paragraph:
            return

        SOFT_CAP = 1200  # internal safety cap (not exposed via CLI)
        if len(paragraph) <= SOFT_CAP:
            out.append(paragraph)
            return

        # Soft split by sentences
        try:
            import pysbd  # type: ignore
            seg = pysbd.Segmenter(language="en", clean=True)
            sents = [s.strip() for s in seg.segment(paragraph) if s.strip()]
        except Exception:
            sents = [s.strip() for s in re.split(r'(?<=[.!?])\s+', paragraph) if s.strip()]

        buf = ""
        for s in sents:
            if not buf:
                buf = s
            elif len(buf) + 1 + len(s) <= SOFT_CAP:
                buf = f"{buf} {s}"
            else:
                out.append(buf)
                buf = s
        if buf:
            out.append(buf)

    # Build paragraphs from units
    for a in atoms:
        if _is_heading(a):
            flush_cur()
            out.append(a)  # heading = standalone chunk
            continue
        if _is_bullet(a):
            flush_cur()
            out.append(a)  # bullet = standalone chunk
            continue
        # Regular sentence → accumulate into current paragraph
        cur.append(a)

    flush_cur()
    return out



def do_split(strategy: str, text: str, fixed_size: int, fixed_overlap: int) -> List[str]:
    """
    Route to the selected splitting strategy.
    'paragraphs' uses document structure only (no CLI size).
    """
    if strategy == "fixed":
        return split_fixed(text, fixed_size, fixed_overlap)
    if strategy == "sentences":
        # One chunk per sentence/heading/bullet
        return split_sentences(text, merge=False)
    if strategy == "paragraphs":
        return split_paragraphs(text)
    raise ValueError("Unknown strategy: choose from fixed/sentences/paragraphs")


# ------------------------------ Gemini Embeddings ------------------------------

def configure_gemini():
    """Load .env and configure the Gemini client."""
    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY missing (.env)")
    genai.configure(api_key=api_key)

def _extract_values_from_resp(resp):
    """Normalize multiple possible google-generativeai response shapes."""
    if isinstance(resp, dict) and "embedding" in resp:
        emb = resp["embedding"]
        if isinstance(emb, dict) and "values" in emb:
            return emb["values"]
        if isinstance(emb, list) and emb and isinstance(emb[0], float):
            return emb
    if isinstance(resp, dict) and "embeddings" in resp:
        embs = resp["embeddings"]
        if isinstance(embs, list) and embs:
            if isinstance(embs[0], dict) and "values" in embs[0]:
                return embs[0]["values"]
            if isinstance(embs[0], list) and embs[0] and isinstance(embs[0][0], float):
                return embs[0]
    if hasattr(resp, "embedding"):
        emb = getattr(resp, "embedding")
        if isinstance(emb, dict) and "values" in emb:
            return emb["values"]
        if isinstance(emb, list) and emb and isinstance(emb[0], float):
            return emb
    if isinstance(resp, list) and resp and isinstance(resp[0], float):
        return resp
    raise RuntimeError(f"Unexpected embedding response format: {type(resp)} -> {resp}")

def embed_texts(texts: List[str], model_name: str = "text-embedding-004", sleep_sec: float = 0.02) -> List[List[float]]:
    """Embed each chunk and return a list of vectors (with optional pacing)."""
    vectors: List[List[float]] = []
    for t in tqdm(texts, desc="Embedding", unit="chunk"):
        resp = genai.embed_content(model=model_name, content=t, task_type="retrieval_document")
        values = _extract_values_from_resp(resp)
        vectors.append(values)
        if sleep_sec:
            time.sleep(sleep_sec)  # spacing to avoid rate limits
    return vectors


# ------------------------------ Postgres ------------------------------

def get_conn():
    """Create a psycopg2 connection using POSTGRES_URL from .env."""
    load_dotenv()
    url = os.getenv("POSTGRES_URL")
    if not url:
        raise RuntimeError("POSTGRES_URL missing (.env)")
    return psycopg2.connect(url)

def ensure_schema(cur, dim: int):
    """Ensure pgvector extension, table, and ANN index exist (idempotent)."""
    cur.execute("CREATE EXTENSION IF NOT EXISTS vector;")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS public.document_chunks (
            id             SERIAL PRIMARY KEY,
            chunk_text     TEXT NOT NULL,
            embedding      VECTOR(%s) NOT NULL,
            filename       TEXT NOT NULL,
            split_strategy TEXT NOT NULL,
            created_at     TIMESTAMPTZ NOT NULL DEFAULT NOW()
        );
    """, (dim,))
    cur.execute("""
        CREATE INDEX IF NOT EXISTS idx_document_chunks_embedding
        ON public.document_chunks
        USING ivfflat (embedding vector_cosine_ops)
        WITH (lists = 100);
    """)

def insert_rows(cur, rows: List[Tuple[str, list, str, str]]):
    """Bulk insert rows: (chunk_text, embedding, filename, split_strategy)."""
    template = "(" + ",".join(["%s"]*4) + ")"
    execute_values(cur,
        "INSERT INTO public.document_chunks (chunk_text, embedding, filename, split_strategy) VALUES %s",
        rows, template=template
    )


# ------------------------------ Main ------------------------------

def main():
    """CLI entry point: read files, split, embed, and store in Postgres."""
    ap = argparse.ArgumentParser()
    ap.add_argument("--paths", nargs="+", required=True, help="Files/directories/globs containing PDF/DOCX")
    ap.add_argument("--strategy", choices=["fixed", "sentences", "paragraphs"], default="fixed")
    ap.add_argument("--chunk-size", type=int, default=800, help="for --strategy fixed or as max_chars when merging")
    ap.add_argument("--overlap", type=int, default=120, help="for --strategy fixed")
    ap.add_argument("--model", default="text-embedding-004")
    args = ap.parse_args()

    configure_gemini()
    files = discover_files(args.paths)
    if not files:
        print("No files found under given --paths.")
        return

    for path in files:
        print(f"Reading: {path}")
        raw = read_any(path)
        text = clean_text(raw)
        chunks = do_split(args.strategy, text, args.chunk_size, args.overlap)
        print(f" → {len(chunks)} chunks")

        vecs = embed_texts(chunks, model_name=args.model)
        if not vecs:
            print("No vectors produced; skipping.")
            continue
        dim = len(vecs[0])

        with get_conn() as conn:
            with conn.cursor() as cur:
                ensure_schema(cur, dim)
                rows = [(c, v, os.path.basename(path), args.strategy) for c, v in zip(chunks, vecs)]
                insert_rows(cur, rows)
                conn.commit()
        print(f"Saved {len(vecs)} chunks for {os.path.basename(path)}")

    print("Done.")

if __name__ == "__main__":
    main()
